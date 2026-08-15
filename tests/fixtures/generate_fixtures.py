"""
Generates the synthetic, deliberately-problematic (and one deliberately
clean) test resumes used by tests/test_extraction.py and
tests/test_ats_engine.py. Run directly to (re)generate fixtures:

    python tests/fixtures/generate_fixtures.py

No real personal data anywhere in this file or its output -- every name,
email, and employer below is fictional. Never replace these with a real
resume; fixtures are committed to git.
"""
from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

FIXTURES_DIR = Path(__file__).resolve().parent
PAGE_W, PAGE_H = letter  # 612 x 792


def _clean_baseline_pdf():
    path = FIXTURES_DIR / "clean_baseline.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, 750, "Jane A. Doe")

    c.setFont("Helvetica", 10)
    c.drawString(50, 730, "Senior Program Manager")
    c.drawString(50, 700, "jane.doe@example.com | (555) 123-4567 | Denver, CO")

    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, 665, "PROFESSIONAL EXPERIENCE")
    c.setFont("Helvetica", 10)
    c.drawString(50, 645, "Program Manager -- Example Aerospace Co.  |  2019-Present")
    c.drawString(50, 628, "- Directed a cross-functional team of 12 to deliver first-article qualification")
    c.drawString(50, 614, "  three months ahead of the contracted schedule.")
    c.drawString(50, 598, "- Owned a $40M program budget across five suppliers with zero cost overruns.")

    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, 565, "EDUCATION")
    c.setFont("Helvetica", 10)
    c.drawString(50, 545, "B.S. Aerospace Engineering, Example State University, 2014")

    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, 515, "SKILLS")
    c.setFont("Helvetica", 10)
    c.drawString(50, 495, "Program management, earned value management, supplier management, risk management")

    c.save()
    return path


def _multi_column_pdf():
    path = FIXTURES_DIR / "two_column.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, 750, "John Q. Smith")
    c.setFont("Helvetica", 10)
    c.drawString(50, 700, "john.smith@example.com | (555) 987-6543")

    # Left column
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, 660, "EXPERIENCE")
    c.setFont("Helvetica", 9)
    c.drawString(50, 640, "Program Manager, Example Co.")
    c.drawString(50, 625, "2018-Present")
    c.drawString(50, 605, "- Led production ramp.")

    # Right column (clear gutter gap from left column's right edge)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(330, 660, "SKILLS")
    c.setFont("Helvetica", 9)
    c.drawString(330, 640, "Risk management")
    c.drawString(330, 625, "Supplier management")
    c.drawString(330, 605, "Earned value mgmt")

    c.save()
    return path


def _image_only_pdf():
    path = FIXTURES_DIR / "image_only.pdf"
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 400, 200))
    pix.set_rect(pix.irect, (200, 200, 200))  # flat gray fill, simulates a scanned page

    doc = fitz.open()
    page = doc.new_page(width=612, height=792)
    page.insert_image(fitz.Rect(50, 50, 562, 742), pixmap=pix)
    doc.save(str(path))
    doc.close()
    return path


def _add_standard_sections(document: Document, include_contact: bool = True):
    if include_contact:
        document.add_paragraph("Alex R. Nguyen")
        document.add_paragraph("alex.nguyen@example.com | (555) 222-3333")
    document.add_heading("Experience", level=1)
    document.add_paragraph("Program Manager, Example Defense Corp., 2017-Present")
    document.add_paragraph("Directed a $25M avionics program from PDR through CDR.", style=None)
    document.add_heading("Education", level=1)
    document.add_paragraph("B.S. Systems Engineering, Example University, 2013")
    document.add_heading("Skills", level=1)
    document.add_paragraph("Program management, supplier management, risk management")


def _table_resume_docx():
    path = FIXTURES_DIR / "table_resume.docx"
    document = Document()
    _add_standard_sections(document)

    document.add_heading("Certifications", level=1)
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    data = [
        ("Certification", "Year"),
        ("PMP", "2019"),
        ("Certified Scrum Master", "2021"),
    ]
    for row, (a, b) in zip(table.rows, data):
        row.cells[0].text = a
        row.cells[1].text = b

    document.save(str(path))
    return path


def _header_footer_contact_docx():
    path = FIXTURES_DIR / "header_footer_contact.docx"
    document = Document()

    section = document.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "Alex R. Nguyen | alex.nguyen@example.com | (555) 222-3333"

    # Deliberately NO contact info in the body -- only in the header.
    _add_standard_sections(document, include_contact=False)

    document.save(str(path))
    return path


def _hidden_text_docx():
    path = FIXTURES_DIR / "hidden_text.docx"
    document = Document()
    _add_standard_sections(document)

    p = document.add_paragraph()
    hidden_run = p.add_run(
        "AWS Azure Kubernetes DevOps Machine Learning Data Science Python Java hidden keyword stuffing"
    )
    hidden_run.font.hidden = True

    p2 = document.add_paragraph()
    white_run = p2.add_run("Six Sigma Black Belt Agile Scrum Lean Manufacturing white text keyword stuffing")
    from docx.shared import RGBColor

    white_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    document.save(str(path))
    return path


def generate_all():
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        _clean_baseline_pdf(),
        _multi_column_pdf(),
        _image_only_pdf(),
        _table_resume_docx(),
        _header_footer_contact_docx(),
        _hidden_text_docx(),
    ]
    for path in created:
        print(f"generated: {path}")


if __name__ == "__main__":
    generate_all()
