"""
DOCX extraction and structural introspection via python-docx + raw
WordprocessingML XML inspection.

python-docx's high-level API (`document.paragraphs`, `document.tables`)
only walks the *normal* document flow. Content placed in text boxes
(`w:txbxContent`) is invisible to that API -- which is exactly the ATS
risk we need to detect (visually present, but frequently unreachable by
an ATS's text extraction). So this module does two passes: the normal
python-docx walk for paragraphs/tables/headers/footers/style names, and a
raw XPath walk for text boxes, hidden runs, white-on-white runs, and
section column counts.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


@dataclass
class DocxParagraph:
    text: str
    style_name: str
    is_heading: bool


@dataclass
class DocxTable:
    n_rows: int
    n_cols: int
    sample_text: str


@dataclass
class DocxExtraction:
    full_text: str
    paragraphs: list[DocxParagraph]
    tables: list[DocxTable]
    header_text: list[str]
    footer_text: list[str]
    column_count: int
    has_text_boxes: bool
    textbox_text: list[str]
    hidden_text_runs: list[str]
    white_text_runs: list[str]
    hyperlink_count: int
    method: str = "python-docx"
    error: str | None = None


def _is_heading_style(style_name: str) -> bool:
    lowered = (style_name or "").lower()
    return lowered.startswith("heading") or lowered == "title"


def extract(path: Path) -> DocxExtraction:
    try:
        document = Document(str(path))
    except Exception as exc:
        return DocxExtraction(
            full_text="",
            paragraphs=[],
            tables=[],
            header_text=[],
            footer_text=[],
            column_count=1,
            has_text_boxes=False,
            textbox_text=[],
            hidden_text_runs=[],
            white_text_runs=[],
            hyperlink_count=0,
            error=str(exc),
        )

    paragraphs: list[DocxParagraph] = []
    for p in document.paragraphs:
        style_name = p.style.name if p.style is not None else ""
        text = p.text.strip()
        if text:
            paragraphs.append(
                DocxParagraph(text=text, style_name=style_name, is_heading=_is_heading_style(style_name))
            )

    tables: list[DocxTable] = []
    for t in document.tables:
        n_rows = len(t.rows)
        n_cols = len(t.columns)
        sample_cells = []
        for row in t.rows[:2]:
            for cell in row.cells:
                if cell.text.strip():
                    sample_cells.append(cell.text.strip())
        tables.append(DocxTable(n_rows=n_rows, n_cols=n_cols, sample_text=" | ".join(sample_cells[:6])))

    header_text: list[str] = []
    footer_text: list[str] = []
    for section in document.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                header_text.append(p.text.strip())
        for p in section.footer.paragraphs:
            if p.text.strip():
                footer_text.append(p.text.strip())

    root = document.element

    # Column count: max across any w:cols/@w:num in section properties.
    column_count = 1
    for cols_el in root.xpath(".//w:sectPr/w:cols"):
        num = cols_el.get(qn("w:num"))
        if num and num.isdigit():
            column_count = max(column_count, int(num))

    # Text boxes: w:txbxContent holds paragraphs invisible to document.paragraphs.
    textbox_els = root.xpath(".//w:txbxContent")
    textbox_text: list[str] = []
    for tb in textbox_els:
        runs_text = "".join(t.text or "" for t in tb.xpath(".//w:t"))
        if runs_text.strip():
            textbox_text.append(runs_text.strip())

    # Hidden runs (w:vanish) and white-on-white runs (w:color w:val="FFFFFF").
    hidden_text_runs: list[str] = []
    white_text_runs: list[str] = []
    for run_el in root.xpath(".//w:r"):
        rpr = run_el.find(qn("w:rPr"))
        run_text = "".join(t.text or "" for t in run_el.findall(qn("w:t")))
        if not run_text.strip():
            continue
        if rpr is None:
            continue
        if rpr.find(qn("w:vanish")) is not None:
            hidden_text_runs.append(run_text.strip())
        color_el = rpr.find(qn("w:color"))
        if color_el is not None:
            val = color_el.get(qn("w:val"))
            if val and val.upper() == "FFFFFF":
                white_text_runs.append(run_text.strip())

    hyperlink_count = len(root.xpath(".//w:hyperlink"))

    full_text = "\n".join(p.text for p in paragraphs)

    return DocxExtraction(
        full_text=full_text,
        paragraphs=paragraphs,
        tables=tables,
        header_text=header_text,
        footer_text=footer_text,
        column_count=column_count,
        has_text_boxes=bool(textbox_text),
        textbox_text=textbox_text,
        hidden_text_runs=hidden_text_runs,
        white_text_runs=white_text_runs,
        hyperlink_count=hyperlink_count,
    )
